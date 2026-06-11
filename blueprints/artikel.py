"""
Blueprint: Artikel Generator
Menghasilkan artikel menggunakan Groq API (LLaMA 3.3 70B),
auto-generate word cloud, dan export ke .docx / .png.
"""

import base64
import glob
import hashlib
import io
import logging
import os
import time
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from flask import Blueprint, jsonify, render_template, request, send_file, session
from groq import APIStatusError, Groq

from utils.image_gen import generate_wordcloud

logger = logging.getLogger(__name__)

# Definisi blueprint
artikel_bp = Blueprint(
    "artikel",
    __name__,
    url_prefix="/artikel",
)

# Word cloud cache TTL — must match app.config["CACHE_DEFAULT_TIMEOUT"]
_WC_CACHE_TTL = 3600  # 1 hour


def _get_cache():
    """Lazy import of the app-level cache to avoid circular imports."""
    from app import cache  # noqa: PLC0415

    return cache


def _cleanup_old_wordclouds(
    cache_dir: str, max_age_seconds: int = _WC_CACHE_TTL
) -> None:
    """
    Remove .png files in cache_dir that are older than max_age_seconds.
    Called on every generate request to prevent disk pile-up.
    """
    now = time.time()
    pattern = os.path.join(cache_dir, "wc_*.png")
    removed = 0
    for fpath in glob.glob(pattern):
        try:
            if now - os.path.getmtime(fpath) > max_age_seconds:
                os.remove(fpath)
                removed += 1
        except OSError:
            pass
    if removed:
        logger.info(
            f"[wordcloud] Cleaned up {removed} expired PNG file(s) from {cache_dir}"
        )


# Inisialisasi Groq client
_groq_client = None


def _get_groq_client():
    """Lazy-init Groq client agar API key dibaca saat pertama kali dipakai."""
    global _groq_client
    if _groq_client is None:
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY belum dikonfigurasi di .env")
        _groq_client = Groq(api_key=api_key)
    return _groq_client


@artikel_bp.route("/")
def index():
    """Halaman utama Artikel Generator."""
    return render_template("artikel.html")


@artikel_bp.route("/generate", methods=["POST"])
def generate():
    """
    Generate artikel menggunakan Groq API (LLaMA 3.3 70B).
    Menerima JSON: { topik, keywords, tone, panjang }
    """
    try:
        data = request.get_json()

        if not data:
            return jsonify({"success": False, "message": "Data tidak valid"}), 400

        topik = data.get("topik", "").strip()
        keywords = data.get("keywords", "").strip()
        tone = data.get("tone", "formal").strip()
        panjang = data.get("panjang", "sedang").strip()

        # Validasi topik
        if not topik:
            return jsonify(
                {"success": False, "message": "Topik artikel wajib diisi"}
            ), 400

        # Mapping panjang ke jumlah kata
        panjang_map = {
            "pendek": "300-500",
            "sedang": "600-900",
            "panjang": "1000-1500",
        }
        jumlah_kata = panjang_map.get(panjang, "600-900")

        # Tentukan tone
        tone_desc = (
            "formal dan profesional"
            if tone == "formal"
            else "santai dan conversational"
        )

        # Konstruksi prompt untuk LLM
        prompt = f"""Buatkan artikel dalam Bahasa Indonesia dengan spesifikasi berikut:

Topik: {topik}
{"Keywords: " + keywords if keywords else ""}
Tone: {tone_desc}
Panjang: {jumlah_kata} kata

Instruksi:
1. Buat artikel yang informatif dan well-structured
2. Gunakan subjudul (heading) untuk setiap bagian
3. Mulai dengan paragraf pembuka yang menarik
4. Akhiri dengan kesimpulan yang kuat
5. {"Sertakan keyword berikut secara natural: " + keywords if keywords else "Buat konten yang relevan dan informatif"}
6. Jangan gunakan format markdown (seperti ** atau ##), gunakan plain text dengan huruf kapital untuk subjudul
7. Pastikan tone penulisan {tone_desc}

Tulis artikelnya langsung tanpa tambahan penjelasan."""

        # Panggil Groq API dengan model LLaMA 3.3 70B
        client = _get_groq_client()
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Kamu adalah penulis artikel profesional untuk PT Telkom Indonesia. "
                        "Tulis artikel yang formal, informatif, dan sesuai tone yang diminta."
                    ),
                },
                {
                    "role": "user",
                    "content": prompt,
                },
            ],
            temperature=0.7,
            max_tokens=2048,
        )

        # Ambil teks hasil generate
        article_text = response.choices[0].message.content
        if not article_text:
            return jsonify(
                {"success": False, "message": "API tidak mengembalikan hasil"}
            ), 500

        article_text = article_text.strip()

        # Generate word cloud — keyed by MD5 of article content (app-level cache)
        article_hash = hashlib.md5(article_text.encode()).hexdigest()
        cache_key = f"wordcloud_{article_hash}"
        app_cache = _get_cache()

        # Run cleanup of stale PNGs on each generate request
        from flask import current_app

        wc_cache_dir = current_app.config.get("CACHE_DIR", "")
        if wc_cache_dir:
            _cleanup_old_wordclouds(wc_cache_dir)

        wordcloud_b64 = app_cache.get(cache_key)
        if wordcloud_b64 is None:
            wordcloud_b64 = generate_wordcloud(article_text)
            if wordcloud_b64:
                app_cache.set(cache_key, wordcloud_b64, timeout=_WC_CACHE_TTL)
            logger.info(f"Word cloud generated and cached (key={cache_key[:12]}...)")
        else:
            logger.info(f"Word cloud cache hit (key={cache_key[:12]}...)")

        # Simpan ke session untuk download nanti
        session["wordcloud_b64"] = wordcloud_b64
        session["article_text"] = article_text
        session["article_topik"] = topik
        session["article_keywords"] = keywords

        # URL word cloud untuk ditampilkan di frontend
        wordcloud_url = (
            f"data:image/png;base64,{wordcloud_b64}" if wordcloud_b64 else None
        )

        logger.info(f"Artikel berhasil di-generate: {topik}")

        return jsonify(
            {
                "success": True,
                "article_text": article_text,
                "wordcloud_url": wordcloud_url,
            }
        )

    except APIStatusError as e:
        # Tangani rate limit (429) secara spesifik
        if e.status_code == 429:
            logger.warning(f"[ERROR] artikel - Groq rate limit tercapai: {e}")
            return jsonify(
                {
                    "success": False,
                    "message": "Rate limit tercapai, coba beberapa saat lagi.",
                }
            ), 429
        logger.error(f"[ERROR] artikel - Groq API error: {e}")
        return jsonify(
            {
                "success": False,
                "message": "Terjadi kesalahan saat menghubungi API. Coba lagi nanti.",
            }
        ), 500

    except ValueError as e:
        # API key belum dikonfigurasi
        logger.error(f"[ERROR] artikel - Konfigurasi error: {e}")
        return jsonify(
            {
                "success": False,
                "message": "Konfigurasi API belum lengkap. Hubungi administrator.",
            }
        ), 500

    except Exception as e:
        logger.error(f"[ERROR] artikel - Error generate artikel: {e}")
        return jsonify(
            {
                "success": False,
                "message": "Terjadi kesalahan server saat generate artikel",
            }
        ), 500


@artikel_bp.route("/download/docx")
def download_docx():
    """
    Download artikel sebagai file Word (.docx).
    """
    try:
        article_text = session.get("article_text")
        topik = session.get("article_topik", "Artikel")

        if not article_text:
            return jsonify(
                {
                    "error": "Tidak ada artikel untuk di-download. Generate terlebih dahulu."
                }
            ), 404

        # Buat dokumen Word
        doc = Document()

        # Style dokumen
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Judul
        title = doc.add_heading(topik, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_run = meta.add_run(
            f"Digenerate pada: {datetime.now().strftime('%d %B %Y, %H:%M')}"
        )
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = None  # default color

        doc.add_paragraph("")  # spacer

        # Konten artikel
        paragraphs = article_text.split("\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue

            # Deteksi apakah ini subjudul (huruf kapital semua atau pendek)
            if para_text.isupper() and len(para_text) < 100:
                doc.add_heading(para_text.title(), level=2)
            else:
                p = doc.add_paragraph(para_text)
                p.paragraph_format.space_after = Pt(6)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph("---")
        footer_text = doc.add_paragraph(
            "Artikel ini digenerate menggunakan AI oleh Productivity Suite — PT Telkom Indonesia"
        )
        footer_run = footer_text.runs[0]
        footer_run.font.size = Pt(8)

        # Simpan ke memory
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        # Bersihkan nama file
        safe_title = "".join(c if c.isalnum() or c in " _-" else "" for c in topik)[:50]
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        logger.error(f"Error download docx: {e}")
        return jsonify({"error": "Gagal membuat file .docx"}), 500


@artikel_bp.route("/download/wordcloud")
def download_wordcloud():
    """
    Download word cloud sebagai file PNG.
    """
    try:
        wordcloud_b64 = session.get("wordcloud_b64")

        if not wordcloud_b64:
            return jsonify(
                {
                    "error": "Tidak ada word cloud untuk di-download. Generate terlebih dahulu."
                }
            ), 404

        # Decode base64 ke bytes
        wordcloud_bytes = base64.b64decode(wordcloud_b64)
        output = io.BytesIO(wordcloud_bytes)
        output.seek(0)

        topik = session.get("article_topik", "wordcloud")
        safe_title = "".join(c if c.isalnum() or c in " _-" else "" for c in topik)[:50]
        filename = f"wordcloud_{safe_title}.png"

        return send_file(
            output,
            mimetype="image/png",
            as_attachment=True,
            download_name=filename,
        )

    except Exception as e:
        logger.error(f"Error download word cloud: {e}")
        return jsonify({"error": "Gagal mengambil word cloud"}), 500
